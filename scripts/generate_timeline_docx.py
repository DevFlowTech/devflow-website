import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Color Palette Constants ---
NAVY_HEX = "0B1220"       # #0B1220 - Dark Navy (Primary Brand)
TEAL_HEX = "14B8A6"       # #14B8A6 - Teal (Primary Accent)
BLUE_HEX = "2563EB"       # #2563EB - Royal Blue (Secondary Accent)
DARK_TEXT_HEX = "0F172A"  # #0F172A - Slate 900
BODY_TEXT_HEX = "334155"  # #334155 - Slate 700
MUTED_HEX = "64748B"      # #64748B - Slate 500
BG_LIGHT_HEX = "F8FAFC"   # #F8FAFC - Light Gray / Slate 50
BG_TEAL_LIGHT = "F0FDFA"  # #F0FDFA - Light Teal Tint
BORDER_HEX = "E2E8F0"     # #E2E8F0 - Subtle Slate Border
BORDER_DARK_HEX = "CBD5E1"# #CBD5E1 - Medium Border
WHITE_HEX = "FFFFFF"      # #FFFFFF - Pure White
GOLD_HEX = "D97706"       # Amber / Highlight

COLOR_NAVY = RGBColor(11, 18, 32)
COLOR_TEAL = RGBColor(20, 184, 166)
COLOR_BLUE = RGBColor(37, 99, 235)
COLOR_DARK = RGBColor(15, 23, 42)
COLOR_BODY = RGBColor(51, 65, 85)
COLOR_MUTED = RGBColor(100, 116, 139)

def set_cell_background(cell, hex_color):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Sets inner margins (padding) of a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Sets specific borders on a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = ['<w:tcBorders ' + nsdecls("w") + '>']
    for side, border_spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if border_spec:
            val = border_spec.get('val', 'single')
            sz = border_spec.get('sz', '4')
            space = border_spec.get('space', '0')
            color = border_spec.get('color', 'auto')
            borders.append(f'<w:{side} w:val="{val}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>')
        else:
            borders.append(f'<w:{side} w:val="none"/>')
    borders.append('</w:tcBorders>')
    tcBorders = parse_xml(''.join(borders))
    tcPr.append(tcBorders)

def add_styled_paragraph(doc, text="", font_name="Segoe UI", font_size=10, bold=False, italic=False, 
                         color=COLOR_BODY, space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15):
    """Helper to add clean styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
    return p

def add_header_footer(doc):
    """Configures professional executive headers and footers with dynamic page numbering."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(4)
    hrun1 = hp.add_run("DEVFLOW TECHNOLOGY  |  ")
    hrun1.font.name = "Segoe UI"
    hrun1.font.size = Pt(8)
    hrun1.font.bold = True
    hrun1.font.color.rgb = COLOR_NAVY

    hrun2 = hp.add_run("ONPOINT GROUP LIMITED — 14-DAY PROJECT TIMELINE")
    hrun2.font.name = "Segoe UI"
    hrun2.font.size = Pt(8)
    hrun2.font.color.rgb = COLOR_MUTED

    # Header bottom border line in XML
    pPr = hp._p.get_or_add_pPr()
    pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="4" w:color="{BORDER_HEX}"/></w:pBdr>')
    pPr.append(pbdr)

    # Footer
    footer = section.footer
    tbl_footer = footer.add_table(rows=1, cols=2, width=Inches(7.1))
    tbl_footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_footer.autofit = False

    cell_l = tbl_footer.rows[0].cells[0]
    cell_r = tbl_footer.rows[0].cells[1]
    cell_l.width = Inches(4.5)
    cell_r.width = Inches(2.6)

    set_cell_margins(cell_l, top=60, bottom=60, left=0, right=60)
    set_cell_margins(cell_r, top=60, bottom=60, left=60, right=0)
    set_cell_borders(cell_l)
    set_cell_borders(cell_r)

    # Footer Left
    fp_l = cell_l.paragraphs[0]
    fp_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp_l.paragraph_format.space_before = Pt(0)
    fp_l.paragraph_format.space_after = Pt(0)
    
    fr1 = fp_l.add_run("DevFlow Technology ")
    fr1.font.name = "Segoe UI"
    fr1.font.size = Pt(8)
    fr1.font.bold = True
    fr1.font.color.rgb = COLOR_NAVY

    fr2 = fp_l.add_run("— Strategic Technology & Digital Solutions Partner\n")
    fr2.font.name = "Segoe UI"
    fr2.font.size = Pt(7.5)
    fr2.font.color.rgb = COLOR_MUTED

    fr3 = fp_l.add_run("https://devflow.co.in")
    fr3.font.name = "Segoe UI"
    fr3.font.size = Pt(7.5)
    fr3.font.bold = True
    fr3.font.color.rgb = COLOR_BLUE

    fr4 = fp_l.add_run("  •  Confidential — Prepared for OnPoint Group Limited")
    fr4.font.name = "Segoe UI"
    fr4.font.size = Pt(7.5)
    fr4.font.color.rgb = COLOR_MUTED

    # Footer Right with Page Number Field
    fp_r = cell_r.paragraphs[0]
    fp_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp_r.paragraph_format.space_before = Pt(0)
    fp_r.paragraph_format.space_after = Pt(0)

    fpr_lbl = fp_r.add_run("Page ")
    fpr_lbl.font.name = "Segoe UI"
    fpr_lbl.font.size = Pt(8)
    fpr_lbl.font.color.rgb = COLOR_MUTED

    # Page field XML
    fldSimple1 = OxmlElement('w:fldSimple')
    fldSimple1.set(qn('w:instr'), 'PAGE')
    fp_r._p.append(fldSimple1)

    fpr_mid = fp_r.add_run(" of ")
    fpr_mid.font.name = "Segoe UI"
    fpr_mid.font.size = Pt(8)
    fpr_mid.font.color.rgb = COLOR_MUTED

    fldSimple2 = OxmlElement('w:fldSimple')
    fldSimple2.set(qn('w:instr'), 'NUMPAGES')
    fp_r._p.append(fldSimple2)

def build_document():
    doc = Document()
    add_header_footer(doc)

    # ==========================================
    # 1. EXECUTIVE COMPANY LETTERHEAD
    # ==========================================
    tbl_letterhead = doc.add_table(rows=1, cols=2)
    tbl_letterhead.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_letterhead.autofit = False

    c_logo = tbl_letterhead.rows[0].cells[0]
    c_meta = tbl_letterhead.rows[0].cells[1]
    c_logo.width = Inches(4.0)
    c_meta.width = Inches(3.1)

    set_cell_margins(c_logo, top=60, bottom=80, left=0, right=100)
    set_cell_margins(c_meta, top=60, bottom=80, left=100, right=0)
    set_cell_borders(c_logo)
    set_cell_borders(c_meta)

    # Left: Logo & Company Name
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(4)

    logo_path = os.path.abspath("public/logo.png")
    if os.path.exists(logo_path):
        run_img = p_logo.add_run()
        run_img.add_picture(logo_path, width=Inches(1.05))
    
    p_comp = c_logo.add_paragraph()
    p_comp.paragraph_format.space_before = Pt(4)
    p_comp.paragraph_format.space_after = Pt(0)
    r_cname = p_comp.add_run("DEVFLOW TECHNOLOGY")
    r_cname.font.name = "Segoe UI"
    r_cname.font.size = Pt(13)
    r_cname.font.bold = True
    r_cname.font.color.rgb = COLOR_NAVY

    p_tag = c_logo.add_paragraph()
    p_tag.paragraph_format.space_before = Pt(1)
    p_tag.paragraph_format.space_after = Pt(0)
    r_tag = p_tag.add_run("Strategic Technology & Digital Solutions Partner")
    r_tag.font.name = "Segoe UI"
    r_tag.font.size = Pt(8.5)
    r_tag.font.color.rgb = COLOR_TEAL
    r_tag.font.bold = True

    # Right: Represented By / Contact Box
    p_meta = c_meta.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(2)

    r_rep_lbl = p_meta.add_run("REPRESENTED BY:\n")
    r_rep_lbl.font.name = "Segoe UI"
    r_rep_lbl.font.size = Pt(7.5)
    r_rep_lbl.font.bold = True
    r_rep_lbl.font.color.rgb = COLOR_MUTED

    r_rep_name = p_meta.add_run("Prince Gajjar\n")
    r_rep_name.font.name = "Segoe UI"
    r_rep_name.font.size = Pt(10)
    r_rep_name.font.bold = True
    r_rep_name.font.color.rgb = COLOR_NAVY

    r_rep_title = p_meta.add_run("Founder & Chief Executive Officer\n")
    r_rep_title.font.name = "Segoe UI"
    r_rep_title.font.size = Pt(8.5)
    r_rep_title.font.color.rgb = COLOR_BODY

    r_rep_web = p_meta.add_run("Website: ")
    r_rep_web.font.name = "Segoe UI"
    r_rep_web.font.size = Pt(8)
    r_rep_web.font.color.rgb = COLOR_MUTED

    r_rep_link = p_meta.add_run("https://devflow.co.in\n")
    r_rep_link.font.name = "Segoe UI"
    r_rep_link.font.size = Pt(8)
    r_rep_link.font.bold = True
    r_rep_link.font.color.rgb = COLOR_BLUE

    r_rep_doc = p_meta.add_run("Document Ref: DFT-OPG-14D-2026")
    r_rep_doc.font.name = "Segoe UI"
    r_rep_doc.font.size = Pt(7.5)
    r_rep_doc.font.color.rgb = COLOR_MUTED

    # Letterhead Accent Divider (Navy + Teal)
    tbl_divider = doc.add_table(rows=1, cols=2)
    tbl_divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_divider.autofit = False
    c_div1 = tbl_divider.rows[0].cells[0]
    c_div2 = tbl_divider.rows[0].cells[1]
    c_div1.width = Inches(5.0)
    c_div2.width = Inches(2.1)
    set_cell_background(c_div1, NAVY_HEX)
    set_cell_background(c_div2, TEAL_HEX)
    set_cell_margins(c_div1, top=15, bottom=15, left=0, right=0)
    set_cell_margins(c_div2, top=15, bottom=15, left=0, right=0)
    set_cell_borders(c_div1)
    set_cell_borders(c_div2)
    c_div1.paragraphs[0].paragraph_format.space_before = Pt(0)
    c_div1.paragraphs[0].paragraph_format.space_after = Pt(0)
    c_div2.paragraphs[0].paragraph_format.space_before = Pt(0)
    c_div2.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Spacing
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(6)
    p_spacer.paragraph_format.space_after = Pt(0)

    # ==========================================
    # 2. DOCUMENT TITLE & PROJECT HERO BANNER
    # ==========================================
    tbl_hero = doc.add_table(rows=1, cols=1)
    tbl_hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_hero.autofit = False
    c_hero = tbl_hero.rows[0].cells[0]
    c_hero.width = Inches(7.1)
    set_cell_background(c_hero, NAVY_HEX)
    set_cell_margins(c_hero, top=200, bottom=200, left=240, right=240)
    set_cell_borders(c_hero)

    p_h0 = c_hero.paragraphs[0]
    p_h0.paragraph_format.space_before = Pt(0)
    p_h0.paragraph_format.space_after = Pt(3)
    r_h0 = p_h0.add_run("OFFICIAL CLIENT PROJECT TIMELINE & DELIVERY SCHEDULE")
    r_h0.font.name = "Segoe UI"
    r_h0.font.size = Pt(8.5)
    r_h0.font.bold = True
    r_h0.font.color.rgb = COLOR_TEAL

    p_h1 = c_hero.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(0)
    p_h1.paragraph_format.space_after = Pt(3)
    r_h1 = p_h1.add_run("ONPOINT GROUP LIMITED")
    r_h1.font.name = "Segoe UI"
    r_h1.font.size = Pt(20)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(255, 255, 255)

    p_h2 = c_hero.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(0)
    p_h2.paragraph_format.space_after = Pt(12)
    r_h2 = p_h2.add_run("Website Redesign & Digital Platform Engineering — 14-Day Delivery Roadmap")
    r_h2.font.name = "Segoe UI"
    r_h2.font.size = Pt(11)
    r_h2.font.color.rgb = RGBColor(226, 232, 240)

    # Sub-table inside hero for metadata cards
    tbl_pills = c_hero.add_table(rows=1, cols=4)
    tbl_pills.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pills.autofit = False
    pill_widths = [Inches(1.65), Inches(1.65), Inches(1.65), Inches(1.65)]
    pill_data = [
        ("PROJECT DOMAIN", "onpointgroup.ng"),
        ("DELIVERY TIMELINE", "14 Calendar Days"),
        ("PREPARED BY", "DevFlow Technology"),
        ("TARGET AUDIENCE", "OnPoint Group (Nigeria)")
    ]
    for i in range(4):
        cp = tbl_pills.rows[0].cells[i]
        cp.width = pill_widths[i]
        set_cell_background(cp, "132036")
        set_cell_margins(cp, top=80, bottom=80, left=80, right=80)
        set_cell_borders(cp, 
                         top={'val': 'single', 'sz': '4', 'color': '223456'},
                         bottom={'val': 'single', 'sz': '4', 'color': '223456'},
                         left={'val': 'single', 'sz': '4', 'color': '223456'},
                         right={'val': 'single', 'sz': '4', 'color': '223456'})
        pp = cp.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(1)
        r_lbl = pp.add_run(pill_data[i][0] + "\n")
        r_lbl.font.name = "Segoe UI"
        r_lbl.font.size = Pt(6.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_TEAL

        r_val = pp.add_run(pill_data[i][1])
        r_val.font.name = "Segoe UI"
        r_val.font.size = Pt(8.5)
        r_val.font.bold = True
        r_val.font.color.rgb = RGBColor(255, 255, 255)

    # ==========================================
    # 3. PROJECT OBJECTIVE & SCOPE OVERVIEW
    # ==========================================
    p_sec1 = doc.add_paragraph()
    p_sec1.paragraph_format.space_before = Pt(14)
    p_sec1.paragraph_format.space_after = Pt(4)
    r_s1_num = p_sec1.add_run("01. ")
    r_s1_num.font.name = "Segoe UI"
    r_s1_num.font.size = Pt(12)
    r_s1_num.font.bold = True
    r_s1_num.font.color.rgb = COLOR_TEAL

    r_s1_txt = p_sec1.add_run("PROJECT OBJECTIVE & EXECUTIVE SUMMARY")
    r_s1_txt.font.name = "Segoe UI"
    r_s1_txt.font.size = Pt(12)
    r_s1_txt.font.bold = True
    r_s1_txt.font.color.rgb = COLOR_NAVY

    # Callout Box for Project Objective
    tbl_obj = doc.add_table(rows=1, cols=1)
    tbl_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_obj.autofit = False
    c_obj = tbl_obj.rows[0].cells[0]
    c_obj.width = Inches(7.1)
    set_cell_background(c_obj, BG_LIGHT_HEX)
    set_cell_margins(c_obj, top=140, bottom=140, left=180, right=180)
    set_cell_borders(c_obj, 
                     left={'val': 'single', 'sz': '24', 'color': TEAL_HEX},
                     top={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                     bottom={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                     right={'val': 'single', 'sz': '4', 'color': BORDER_HEX})

    p_obj = c_obj.paragraphs[0]
    p_obj.paragraph_format.space_before = Pt(0)
    p_obj.paragraph_format.space_after = Pt(6)
    p_obj.paragraph_format.line_spacing = 1.2
    r_obj_lead = p_obj.add_run(
        "DevFlow Technology will redesign and redevelop the OnPoint Group Limited website into a modern, "
        "premium, responsive, high-performance digital platform that clearly communicates OnPoint's business capabilities, "
        "strengthens its corporate positioning, improves user experience, and supports future business development."
    )
    r_obj_lead.font.name = "Segoe UI"
    r_obj_lead.font.size = Pt(9.5)
    r_obj_lead.font.bold = True
    r_obj_lead.font.color.rgb = COLOR_DARK

    p_obj2 = c_obj.add_paragraph()
    p_obj2.paragraph_format.space_before = Pt(0)
    p_obj2.paragraph_format.space_after = Pt(0)
    p_obj2.paragraph_format.line_spacing = 1.2
    r_obj_body = p_obj2.add_run(
        "The project includes full UI/UX redesign, responsive frontend development, content integration, technical SEO architecture, "
        "performance optimization, cross-browser quality assurance, and zero-downtime production deployment on the primary domain "
    )
    r_obj_body.font.name = "Segoe UI"
    r_obj_body.font.size = Pt(9)
    r_obj_body.font.color.rgb = COLOR_BODY

    r_obj_dom = p_obj2.add_run("onpointgroup.ng.")
    r_obj_dom.font.name = "Segoe UI"
    r_obj_dom.font.size = Pt(9)
    r_obj_dom.font.bold = True
    r_obj_dom.font.color.rgb = COLOR_BLUE

    # ==========================================
    # 4. 14-DAY MASTER DELIVERY TIMELINE (DETAILED)
    # ==========================================
    p_sec2 = doc.add_paragraph()
    p_sec2.paragraph_format.space_before = Pt(16)
    p_sec2.paragraph_format.space_after = Pt(4)
    r_s2_num = p_sec2.add_run("02. ")
    r_s2_num.font.name = "Segoe UI"
    r_s2_num.font.size = Pt(12)
    r_s2_num.font.bold = True
    r_s2_num.font.color.rgb = COLOR_TEAL

    r_s2_txt = p_sec2.add_run("14-DAY CALENDAR DELIVERY TIMELINE (DAY 1 TO DAY 14)")
    r_s2_txt.font.name = "Segoe UI"
    r_s2_txt.font.size = Pt(12)
    r_s2_txt.font.bold = True
    r_s2_txt.font.color.rgb = COLOR_NAVY

    p_s2_sub = doc.add_paragraph()
    p_s2_sub.paragraph_format.space_before = Pt(0)
    p_s2_sub.paragraph_format.space_after = Pt(8)
    r_s2_sub = p_s2_sub.add_run(
        "Each calendar day has an explicitly assigned technical objective, task breakdown, and verified deliverable to ensure uncompromising project velocity and quality control."
    )
    r_s2_sub.font.name = "Segoe UI"
    r_s2_sub.font.size = Pt(8.5)
    r_s2_sub.font.color.rgb = COLOR_MUTED

    # Timeline Days Data Structure
    timeline_days = [
        {
            "day": "DAY 01",
            "phase": "PHASE 1: STRATEGY & DISCOVERY",
            "title": "Discovery & Project Setup",
            "activities": [
                "Conduct final project briefing and alignment with OnPoint stakeholders",
                "Perform in-depth audit of existing OnPoint website assets and legacy content",
                "Confirm core business objectives, brand voice, and digital positioning goals",
                "Finalize comprehensive sitemap and hierarchical page structure",
                "Review available company information, background history, and executive bios",
                "Confirm primary services, subsidiary capabilities, and industry verticals",
                "Identify required third-party integrations, tracking scripts, and contact endpoints",
                "Establish dedicated communication channels and feedback turnaround protocols",
                "Initialize repository, CI/CD pipeline, and staging development environment"
            ],
            "deliverable": "Approved project structure, comprehensive sitemap, and master development plan."
        },
        {
            "day": "DAY 02",
            "phase": "PHASE 1: STRATEGY & DISCOVERY",
            "title": "Information Architecture & UX Planning",
            "activities": [
                "Finalize global navigation architecture and header/footer taxonomy",
                "Define homepage visual hierarchy and high-impact hero value proposition",
                "Structure individual service page templates and capability breakdowns",
                "Structure industry-specific landing sections and market presentation",
                "Define About / Corporate Company profile structure and leadership presentation",
                "Design frictionless Contact flow and conversion-focused Call-to-Action (CTA) paths",
                "Map end-to-end user journeys for prospective clients, partners, and stakeholders",
                "Establish responsive layout grid, typographic scale, and container standards"
            ],
            "deliverable": "Approved website information architecture, wireframe flow, and UX structure."
        },
        {
            "day": "DAY 03",
            "phase": "PHASE 2: UI/UX DESIGN SYSTEM",
            "title": "Homepage UI/UX Design",
            "activities": [
                "Design high-impact Hero section with bold editorial typography and primary CTA",
                "Design modern sticky navigation with corporate brand badge and contact trigger",
                "Craft Company Positioning and Executive Overview narrative layout",
                "Design Services Overview modular grid with interactive hover paradigms",
                "Create dynamic Business Statistics and impact metrics display",
                "Design Industry Verticals showcase and market capability layout",
                "Design Strategic Technology Partner feature section highlighting DevFlow partnership",
                "Design Corporate Testimonials, Client Endorsements, and Case Study teasers",
                "Design high-conversion Final CTA banner and multi-column corporate Footer"
            ],
            "deliverable": "Complete, high-fidelity Homepage UI/UX design mockup and visual layout."
        },
        {
            "day": "DAY 04",
            "phase": "PHASE 2: UI/UX DESIGN SYSTEM",
            "title": "Inner Page UI/UX Design & Design System Finalization",
            "activities": [
                "Design master Services Hub page and individual specialized service layouts",
                "Design corporate About page (Vision, Mission, Core Values, Leadership)",
                "Design dedicated Industry Verticals and market solutions pages",
                "Design specialized Real Estate portfolio and project presentation section",
                "Design Technology Partner & Strategic Alliances ecosystem showcase",
                "Design interactive Contact page featuring inquiry form, location, and direct channels",
                "Design complete tablet and mobile adaptive viewport layouts",
                "Finalize master Design System (color tokens, font pairings, button states, spacing rules)"
            ],
            "deliverable": "Complete multi-page UI/UX design system and approved responsive interface mockups."
        },
        {
            "day": "DAY 05",
            "phase": "PHASE 3: FRONTEND ENGINEERING",
            "title": "Frontend Development Begins & Core Framework",
            "activities": [
                "Initialize production codebase with modern Next.js / TypeScript component architecture",
                "Configure global design tokens, CSS variables, and bespoke utility system",
                "Build responsive Global Header, dynamic mobile slide-out drawer, and desktop navigation",
                "Build comprehensive multi-column Global Footer with legal links and contact badges",
                "Configure typography scales, font preloading, and iconography library",
                "Establish reusable atomic UI component library (Buttons, Badges, Cards, Modals)",
                "Implement responsive container system and layout wrappers across viewports"
            ],
            "deliverable": "Fully configured core website framework, global navigation, and reusable UI library."
        },
        {
            "day": "DAY 06",
            "phase": "PHASE 3: FRONTEND ENGINEERING",
            "title": "Homepage Development & Micro-Interactions",
            "activities": [
                "Implement high-impact Hero section with smooth entrance animations and action buttons",
                "Code interactive Services Grid with seamless hover transitions and routing",
                "Implement dynamic Statistics counter module and credibility indicators",
                "Build Industries Showcase with tabbed / modular responsive behavior",
                "Code Strategic Technology Partner section showcasing DevFlow collaboration",
                "Build Testimonials carousel and Case Study highlight components",
                "Implement high-conversion CTA banners and inquiry access points",
                "Enforce strict mobile and tablet touch responsiveness across all homepage sections"
            ],
            "deliverable": "Fully functional, responsive, interactive Homepage on staging environment."
        },
        {
            "day": "DAY 07",
            "phase": "PHASE 3: FRONTEND ENGINEERING",
            "title": "Services & Industries Development",
            "activities": [
                "Develop master Services directory and categorized capability filters",
                "Build detailed individual service page templates with structured content modules",
                "Develop comprehensive Industry Verticals presentation and sector overviews",
                "Build dedicated Real Estate showcase section with project spotlight layouts",
                "Implement strategic Business Development modules and B2B engagement triggers",
                "Enforce full responsive fidelity across desktop, tablet, and mobile breakpoints"
            ],
            "deliverable": "Complete, fully responsive Services, Industries, and Real Estate section implementation."
        },
        {
            "day": "DAY 08",
            "phase": "PHASE 3: FRONTEND ENGINEERING",
            "title": "About, Partnership & Contact Development",
            "activities": [
                "Build About OnPoint page (Company Story, Executive Vision, Core Pillars)",
                "Develop Strategic Technology Partner page featuring DevFlow Technology capabilities",
                "Implement Partnership Information and strategic collaboration frameworks",
                "Develop interactive Contact page with multi-field validation and inquiry routing",
                "Integrate asynchronous contact form handler with spam protection and email notifications",
                "Implement direct WhatsApp/Phone call triggers, interactive location map, and social channels"
            ],
            "deliverable": "Completed About, Partnership, and functional Contact inner pages on staging."
        },
        {
            "day": "DAY 09",
            "phase": "PHASE 3: FRONTEND ENGINEERING",
            "title": "Content Integration & Feature Completeness",
            "activities": [
                "Execute comprehensive integration of all finalized client text, copy, and messaging",
                "Integrate high-resolution imagery, brand marks, and multimedia assets",
                "Verify and test end-to-end contact form submissions and notification delivery",
                "Perform full-site navigation verification and cross-page anchor routing check",
                "Audit and validate all internal hyperlinks, external links, and document downloads",
                "Test all button actions, modal triggers, and interactive states",
                "Integrate required third-party analytics and live chat widgets where applicable"
            ],
            "deliverable": "Feature-complete, content-integrated staging website ready for technical audit."
        },
        {
            "day": "DAY 10",
            "phase": "PHASE 4: TECHNICAL SEO & OPTIMIZATION",
            "title": "Technical SEO & Search Architecture",
            "activities": [
                "Formulate and inject optimized Meta Titles and Meta Descriptions for every page",
                "Structure strict semantic heading hierarchy (H1, H2, H3, H4) across the site",
                "Assign descriptive, keyword-rich Alt Text to all images and media elements",
                "Construct strategic internal linking graph to maximize crawl depth and page authority",
                "Generate automated dynamic XML Sitemap (`/sitemap.xml`) and production `robots.txt`",
                "Enforce self-referential Canonical URLs to prevent duplicate indexing issues",
                "Configure Open Graph (OG) and Twitter Card social sharing metadata",
                "Inject JSON-LD Structured Data Schema (`Organization`, `WebSite`, `LocalBusiness`, `Service`)",
                "Format content blocks for generative AI search engines (AEO/GEO readiness)"
            ],
            "deliverable": "Fully configured, SEO-ready and Schema-structured digital platform."
        },
        {
            "day": "DAY 11",
            "phase": "PHASE 4: TECHNICAL SEO & OPTIMIZATION",
            "title": "Performance Optimization & Core Web Vitals",
            "activities": [
                "Compress, convert, and optimize all images into next-generation WebP/AVIF formats",
                "Configure responsive image sizing (`srcset`), priority hero loading, and lazy loading",
                "Execute code minification, bundle tree-shaking, and dead-code elimination",
                "Optimize Google Font delivery with font-display swapping and local caching",
                "Optimize Core Web Vitals metrics (LCP < 2.5s, FID/INP < 100ms, CLS < 0.1)",
                "Configure aggressive browser asset caching and compression headers",
                "Conduct mobile bandwidth throttling tests (Fast 3G/4G simulation) to guarantee instant loads"
            ],
            "deliverable": "High-speed, performance-optimized website achieving 90+ PageSpeed benchmarks."
        },
        {
            "day": "DAY 12",
            "phase": "PHASE 5: QUALITY ASSURANCE & TESTING",
            "title": "Comprehensive Quality Assurance & Cross-Platform Testing",
            "activities": [
                "Execute cross-device QA testing across Desktop (1920px, 1440px), Tablet (768px, 1024px), and Mobile (375px, 414px)",
                "Conduct cross-browser testing across Google Chrome, Apple Safari, Mozilla Firefox, and Microsoft Edge",
                "Validate all form fields, input masks, validation errors, and success state confirmations",
                "Audit all interactive UI elements, dropdowns, drawer menus, accordions, and hover animations",
                "Perform broken link checker scan across all internal and outbound hyperlinks",
                "Verify font rendering, typography scaling, image clarity, and visual alignment",
                "Conduct basic security review (SSL enforcement, form sanitization, header policies)"
            ],
            "deliverable": "QA-approved release candidate ready for official client review and walkthrough."
        },
        {
            "day": "DAY 13",
            "phase": "PHASE 5: QUALITY ASSURANCE & TESTING",
            "title": "Client Review, Walkthrough & Final Revisions",
            "activities": [
                "Conduct live guided client walkthrough and presentation of the staging website",
                "Demonstrate all responsive layouts, contact flows, and service presentations",
                "Collect consolidated client feedback through structured review form",
                "Apply approved minor refinements, text adjustments, and visual tweaks",
                "Perform final content verification and sign-off check with OnPoint leadership",
                "Execute final technical sanity checks on all modified components",
                "Obtain official client sign-off and green light for production launch"
            ],
            "deliverable": "Final client-approved website ready for live production deployment.",
            "important_note": "CRITICAL REQUIREMENT: Client feedback must be consolidated and submitted within the agreed review window on Day 13 to maintain the scheduled Day 14 production launch."
        },
        {
            "day": "DAY 14",
            "phase": "PHASE 6: PRODUCTION LAUNCH & HANDOVER",
            "title": "Production Deployment, DNS Cutover & Official Handover",
            "activities": [
                "Execute zero-downtime production build and deploy to high-availability global CDN",
                "Configure DNS records (A/CNAME/TXT) and verify domain cutover for `onpointgroup.ng`",
                "Provision and verify automated SSL/TLS security certificates for HTTPS encryption",
                "Verify live production XML Sitemap, robots.txt routing, and canonical domains",
                "Connect and verify Google Analytics 4 (GA4) and Google Search Console indexing",
                "Conduct live production smoke testing (contact forms, links, responsive layouts)",
                "Generate full disaster recovery backup and security snapshot",
                "Deliver project handover documentation, credential vault, and admin guidelines"
            ],
            "deliverable": "ONPOINTGROUP.NG — LIVE PRODUCTION WEBSITE & OFFICIAL HANDOVER."
        }
    ]

    # Render each day as a clean, executive card table
    for item in timeline_days:
        tbl_day = doc.add_table(rows=1, cols=1)
        tbl_day.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_day.autofit = False
        c_day = tbl_day.rows[0].cells[0]
        c_day.width = Inches(7.1)
        
        # Left border colored by phase
        left_border_color = TEAL_HEX if "PHASE 6" in item['phase'] or "PHASE 2" in item['phase'] else (
            NAVY_HEX if "PHASE 3" in item['phase'] else (
                GOLD_HEX if "PHASE 4" in item['phase'] else BLUE_HEX
            )
        )
        
        set_cell_background(c_day, BG_LIGHT_HEX)
        set_cell_margins(c_day, top=100, bottom=100, left=140, right=140)
        set_cell_borders(c_day,
                         left={'val': 'single', 'sz': '20', 'color': left_border_color},
                         top={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                         bottom={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                         right={'val': 'single', 'sz': '4', 'color': BORDER_HEX})
        
        # Top Header line of the Card: Day Badge + Phase Pill + Title
        p_dhdr = c_day.paragraphs[0]
        p_dhdr.paragraph_format.space_before = Pt(0)
        p_dhdr.paragraph_format.space_after = Pt(2)
        
        r_dbadge = p_dhdr.add_run(f"[{item['day']}]  ")
        r_dbadge.font.name = "Segoe UI"
        r_dbadge.font.size = Pt(10)
        r_dbadge.font.bold = True
        r_dbadge.font.color.rgb = COLOR_NAVY

        r_dtitle = p_dhdr.add_run(item['title'])
        r_dtitle.font.name = "Segoe UI"
        r_dtitle.font.size = Pt(10.5)
        r_dtitle.font.bold = True
        r_dtitle.font.color.rgb = COLOR_DARK

        p_dphase = c_day.add_paragraph()
        p_dphase.paragraph_format.space_before = Pt(0)
        p_dphase.paragraph_format.space_after = Pt(6)
        r_dphase = p_dphase.add_run(item['phase'].upper())
        r_dphase.font.name = "Segoe UI"
        r_dphase.font.size = Pt(7.5)
        r_dphase.font.bold = True
        r_dphase.font.color.rgb = COLOR_TEAL

        # Activities list
        p_act_lbl = c_day.add_paragraph()
        p_act_lbl.paragraph_format.space_before = Pt(0)
        p_act_lbl.paragraph_format.space_after = Pt(2)
        r_act_lbl = p_act_lbl.add_run("Key Activities & Scope of Work:")
        r_act_lbl.font.name = "Segoe UI"
        r_act_lbl.font.size = Pt(8)
        r_act_lbl.font.bold = True
        r_act_lbl.font.color.rgb = COLOR_MUTED

        for act in item['activities']:
            p_act = c_day.add_paragraph()
            p_act.paragraph_format.space_before = Pt(0)
            p_act.paragraph_format.space_after = Pt(1.5)
            p_act.paragraph_format.left_indent = Inches(0.15)
            
            r_bullet = p_act.add_run("•  ")
            r_bullet.font.name = "Segoe UI"
            r_bullet.font.size = Pt(8.5)
            r_bullet.font.bold = True
            r_bullet.font.color.rgb = COLOR_TEAL

            r_act_txt = p_act.add_run(act)
            r_act_txt.font.name = "Segoe UI"
            r_act_txt.font.size = Pt(8.5)
            r_act_txt.font.color.rgb = COLOR_BODY

        # Deliverable Box inside the Day card
        tbl_deliv = c_day.add_table(rows=1, cols=1)
        tbl_deliv.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_deliv.autofit = False
        c_deliv = tbl_deliv.rows[0].cells[0]
        c_deliv.width = Inches(6.7)
        
        is_final_day = (item['day'] == "DAY 14")
        deliv_bg = "0B1220" if is_final_day else BG_TEAL_LIGHT
        deliv_border = TEAL_HEX if is_final_day else "99F6E4"

        set_cell_background(c_deliv, deliv_bg)
        set_cell_margins(c_deliv, top=60, bottom=60, left=100, right=100)
        set_cell_borders(c_deliv,
                         left={'val': 'single', 'sz': '8', 'color': TEAL_HEX},
                         top={'val': 'single', 'sz': '4', 'color': deliv_border},
                         bottom={'val': 'single', 'sz': '4', 'color': deliv_border},
                         right={'val': 'single', 'sz': '4', 'color': deliv_border})

        p_del = c_deliv.paragraphs[0]
        p_del.paragraph_format.space_before = Pt(0)
        p_del.paragraph_format.space_after = Pt(0)
        
        r_del_lbl = p_del.add_run("KEY DELIVERABLE:  ")
        r_del_lbl.font.name = "Segoe UI"
        r_del_lbl.font.size = Pt(8)
        r_del_lbl.font.bold = True
        r_del_lbl.font.color.rgb = COLOR_TEAL if is_final_day else COLOR_NAVY

        r_del_val = p_del.add_run(item['deliverable'])
        r_del_val.font.name = "Segoe UI"
        r_del_val.font.size = Pt(8.5)
        r_del_val.font.bold = True
        r_del_val.font.color.rgb = RGBColor(255, 255, 255) if is_final_day else COLOR_DARK

        if "important_note" in item:
            p_note = c_day.add_paragraph()
            p_note.paragraph_format.space_before = Pt(4)
            p_note.paragraph_format.space_after = Pt(0)
            r_note = p_note.add_run("⚠️ " + item['important_note'])
            r_note.font.name = "Segoe UI"
            r_note.font.size = Pt(8)
            r_note.font.bold = True
            r_note.font.color.rgb = RGBColor(180, 83, 9)

        # Gap between cards
        p_gap = doc.add_paragraph()
        p_gap.paragraph_format.space_before = Pt(4)
        p_gap.paragraph_format.space_after = Pt(0)

    # ==========================================
    # 5. PROJECT MILESTONE SUMMARY TABLE
    # ==========================================
    p_sec3 = doc.add_paragraph()
    p_sec3.paragraph_format.space_before = Pt(14)
    p_sec3.paragraph_format.space_after = Pt(4)
    r_s3_num = p_sec3.add_run("03. ")
    r_s3_num.font.name = "Segoe UI"
    r_s3_num.font.size = Pt(12)
    r_s3_num.font.bold = True
    r_s3_num.font.color.rgb = COLOR_TEAL

    r_s3_txt = p_sec3.add_run("PROJECT MILESTONE SUMMARY")
    r_s3_txt.font.name = "Segoe UI"
    r_s3_txt.font.size = Pt(12)
    r_s3_txt.font.bold = True
    r_s3_txt.font.color.rgb = COLOR_NAVY

    milestones_data = [
        ("Discovery & Planning", "Day 1", "Approved project structure, sitemap, and development plan"),
        ("Information Architecture", "Day 2", "Approved website UX hierarchy, wireframes & user journeys"),
        ("Homepage UI/UX", "Day 3", "High-fidelity homepage design mockup"),
        ("Complete UI/UX", "Day 4", "Complete multi-page UI/UX design system & responsive views"),
        ("Development Begins", "Day 5", "Core Next.js website architecture, global header & footer"),
        ("Homepage Complete", "Day 6", "Fully functional, responsive homepage on staging"),
        ("Services & Industries", "Day 7", "Complete Services, Industries & Real Estate sections"),
        ("Core Pages Complete", "Day 8", "About, Strategic Technology Partner & Contact pages"),
        ("Feature Complete", "Day 9", "Full content integration, working forms & interactive features"),
        ("SEO Implementation", "Day 10", "Meta tags, structured schema, sitemap, and AEO/GEO indexing"),
        ("Performance Optimization", "Day 11", "Asset compression, Next.js optimization & 90+ PageSpeed"),
        ("QA Complete", "Day 12", "Full cross-browser, cross-device & mobile QA approval"),
        ("Client Approval", "Day 13", "Client walkthrough, consolidated feedback & final sign-off"),
        ("Production Launch", "Day 14", "onpointgroup.ng LIVE on production domain & official handover")
    ]

    tbl_ms = doc.add_table(rows=len(milestones_data) + 1, cols=3)
    tbl_ms.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_ms.autofit = False

    col_widths = [Inches(2.0), Inches(0.9), Inches(4.2)]
    headers = ["Milestone", "Target", "Key Deliverable & Outcome"]

    # Table Header Row
    for j in range(3):
        cell = tbl_ms.rows[0].cells[j]
        cell.width = col_widths[j]
        set_cell_background(cell, NAVY_HEX)
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        set_cell_borders(cell,
                         top={'val': 'single', 'sz': '4', 'color': NAVY_HEX},
                         bottom={'val': 'single', 'sz': '6', 'color': TEAL_HEX},
                         left={'val': 'none'}, right={'val': 'none'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(headers[j])
        run.font.name = "Segoe UI"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Table Data Rows
    for i, row in enumerate(milestones_data):
        row_cells = tbl_ms.rows[i + 1].cells
        bg_row = BG_LIGHT_HEX if i % 2 == 1 else WHITE_HEX
        is_launch = (i == len(milestones_data) - 1)
        if is_launch:
            bg_row = BG_TEAL_LIGHT

        for j in range(3):
            cell = row_cells[j]
            cell.width = col_widths[j]
            set_cell_background(cell, bg_row)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            set_cell_borders(cell,
                             top={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                             bottom={'val': 'single', 'sz': '4', 'color': (TEAL_HEX if is_launch else BORDER_HEX)},
                             left={'val': 'none'}, right={'val': 'none'})
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.add_run(row[j])
            run.font.name = "Segoe UI"
            run.font.size = Pt(8.5)
            if j == 0 or is_launch:
                run.font.bold = True
                run.font.color.rgb = COLOR_NAVY if not is_launch else COLOR_BLUE
            elif j == 1:
                run.font.bold = True
                run.font.color.rgb = COLOR_TEAL if not is_launch else COLOR_BLUE
            else:
                run.font.color.rgb = COLOR_BODY

    # ==========================================
    # 6. CLIENT REQUIREMENTS & CRITICAL DEPENDENCIES
    # ==========================================
    p_sec4 = doc.add_paragraph()
    p_sec4.paragraph_format.space_before = Pt(16)
    p_sec4.paragraph_format.space_after = Pt(4)
    r_s4_num = p_sec4.add_run("04. ")
    r_s4_num.font.name = "Segoe UI"
    r_s4_num.font.size = Pt(12)
    r_s4_num.font.bold = True
    r_s4_num.font.color.rgb = COLOR_TEAL

    r_s4_txt = p_sec4.add_run("CLIENT REQUIREMENTS & PROJECT DEPENDENCIES")
    r_s4_txt.font.name = "Segoe UI"
    r_s4_txt.font.size = Pt(12)
    r_s4_txt.font.bold = True
    r_s4_txt.font.color.rgb = COLOR_NAVY

    tbl_deps = doc.add_table(rows=1, cols=1)
    tbl_deps.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_deps.autofit = False
    c_deps = tbl_deps.rows[0].cells[0]
    c_deps.width = Inches(7.1)
    set_cell_background(c_deps, BG_LIGHT_HEX)
    set_cell_margins(c_deps, top=120, bottom=120, left=160, right=160)
    set_cell_borders(c_deps,
                     left={'val': 'single', 'sz': '20', 'color': BLUE_HEX},
                     top={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                     bottom={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                     right={'val': 'single', 'sz': '4', 'color': BORDER_HEX})

    p_dintro = c_deps.paragraphs[0]
    p_dintro.paragraph_format.space_before = Pt(0)
    p_dintro.paragraph_format.space_after = Pt(6)
    r_dintro = p_dintro.add_run(
        "To achieve the aggressive 14-day production launch, timely client collaboration is essential. "
        "OnPoint Group Limited should provide the following assets and approvals according to the milestone schedule:"
    )
    r_dintro.font.name = "Segoe UI"
    r_dintro.font.size = Pt(8.5)
    r_dintro.font.color.rgb = COLOR_DARK

    client_items = [
        "Final Company Profile, Mission, Vision, and Leadership Biographies",
        "Final Service & Industry capability descriptions",
        "Company History, Milestones, and Business Achievements",
        "Official Corporate Contact Information & Designated Inquiry Email Endpoints",
        "High-Resolution Vector Logo files (SVG, PNG) and Brand Guidelines (if any)",
        "High-Quality Corporate, Project, and Team Photography/Media Assets",
        "Client Testimonials, Partner Endorsements, and Real Estate Project Case Studies",
        "Corporate Social Media profile links (LinkedIn, X/Twitter, Instagram, etc.)",
        "Legal & Compliance content (Privacy Policy, Terms of Use, Regulatory Disclaimers)",
        "Required Third-Party Credentials, Form Endpoints, or Tracking API keys",
        "Domain Registrar & DNS Access credentials for `onpointgroup.ng` cutover (by Day 13)",
        "Hosting / Cloud Server access where applicable",
        "Google Analytics 4 & Google Search Console access permissions where applicable"
    ]

    for c_item in client_items:
        p_ci = c_deps.add_paragraph()
        p_ci.paragraph_format.space_before = Pt(0)
        p_ci.paragraph_format.space_after = Pt(1.5)
        p_ci.paragraph_format.left_indent = Inches(0.15)
        
        r_cb = p_ci.add_run("☑  ")
        r_cb.font.name = "Segoe UI"
        r_cb.font.size = Pt(8.5)
        r_cb.font.bold = True
        r_cb.font.color.rgb = COLOR_BLUE

        r_ct = p_ci.add_run(c_item)
        r_ct.font.name = "Segoe UI"
        r_ct.font.size = Pt(8.5)
        r_ct.font.color.rgb = COLOR_BODY

    # ==========================================
    # 7. GOVERNANCE: FEEDBACK, SCOPE & CONDITIONS
    # ==========================================
    p_sec5 = doc.add_paragraph()
    p_sec5.paragraph_format.space_before = Pt(16)
    p_sec5.paragraph_format.space_after = Pt(4)
    r_s5_num = p_sec5.add_run("05. ")
    r_s5_num.font.name = "Segoe UI"
    r_s5_num.font.size = Pt(12)
    r_s5_num.font.bold = True
    r_s5_num.font.color.rgb = COLOR_TEAL

    r_s5_txt = p_sec5.add_run("PROJECT GOVERNANCE & EXECUTION TERMS")
    r_s5_txt.font.name = "Segoe UI"
    r_s5_txt.font.size = Pt(12)
    r_s5_txt.font.bold = True
    r_s5_txt.font.color.rgb = COLOR_NAVY

    # Sub-section A: Feedback & Approval Process
    p_sub_a = doc.add_paragraph()
    p_sub_a.paragraph_format.space_before = Pt(4)
    p_sub_a.paragraph_format.space_after = Pt(2)
    r_sa = p_sub_a.add_run("A. Feedback & Approval Process")
    r_sa.font.name = "Segoe UI"
    r_sa.font.size = Pt(10)
    r_sa.font.bold = True
    r_sa.font.color.rgb = COLOR_NAVY

    p_fa_desc = doc.add_paragraph()
    p_fa_desc.paragraph_format.space_before = Pt(0)
    p_fa_desc.paragraph_format.space_after = Pt(4)
    p_fa_desc.paragraph_format.line_spacing = 1.15
    r_fa_desc = p_fa_desc.add_run(
        "To maintain the 14-day delivery schedule, DevFlow Technology will work through defined sequential milestones. "
        "Client feedback should be strictly:"
    )
    r_fa_desc.font.name = "Segoe UI"
    r_fa_desc.font.size = Pt(8.5)
    r_fa_desc.font.color.rgb = COLOR_BODY

    fb_points = [
        ("Consolidated: ", "Provided as a single unified feedback document from the OnPoint project lead rather than fragmented communications."),
        ("Clear & Specific: ", "Containing actionable instructions, exact copy replacements, and specific UI references."),
        ("Timely: ", "Delivered strictly within the agreed 24-hour review window on milestone review days (Day 4 and Day 13)."),
        ("Scope-Aligned: ", "Limited strictly to the approved project scope and architectural blueprint.")
    ]
    for b_title, b_desc in fb_points:
        p_fp = doc.add_paragraph()
        p_fp.paragraph_format.space_before = Pt(0)
        p_fp.paragraph_format.space_after = Pt(1.5)
        p_fp.paragraph_format.left_indent = Inches(0.15)
        r_fpt = p_fp.add_run("•  " + b_title)
        r_fpt.font.name = "Segoe UI"
        r_fpt.font.size = Pt(8.5)
        r_fpt.font.bold = True
        r_fpt.font.color.rgb = COLOR_NAVY
        r_fpd = p_fp.add_run(b_desc)
        r_fpd.font.name = "Segoe UI"
        r_fpd.font.size = Pt(8.5)
        r_fpd.font.color.rgb = COLOR_BODY

    # Sub-section B: Scope Control
    p_sub_b = doc.add_paragraph()
    p_sub_b.paragraph_format.space_before = Pt(6)
    p_sub_b.paragraph_format.space_after = Pt(2)
    r_sb = p_sub_b.add_run("B. Scope Control & Change Management")
    r_sb.font.name = "Segoe UI"
    r_sb.font.size = Pt(10)
    r_sb.font.bold = True
    r_sb.font.color.rgb = COLOR_NAVY

    p_sc_desc = doc.add_paragraph()
    p_sc_desc.paragraph_format.space_before = Pt(0)
    p_sc_desc.paragraph_format.space_after = Pt(4)
    p_sc_desc.paragraph_format.line_spacing = 1.15
    r_sc_desc = p_sc_desc.add_run(
        "The above 14-day timeline is based on the currently agreed website scope. "
        "Additional pages, major architectural changes, custom software integrations, substantial copy rewrites, or features introduced "
        "after formal milestone sign-off may require additional development time and/or commercial adjustment. "
        "Minor refinements within the approved scope will be seamlessly incorporated during the designated review stages."
    )
    r_sc_desc.font.name = "Segoe UI"
    r_sc_desc.font.size = Pt(8.5)
    r_sc_desc.font.color.rgb = COLOR_BODY

    # Sub-section C: Important Project Condition (Prominent Callout Box)
    p_sub_c = doc.add_paragraph()
    p_sub_c.paragraph_format.space_before = Pt(6)
    p_sub_c.paragraph_format.space_after = Pt(2)
    r_sc = p_sub_c.add_run("C. Critical Project Condition")
    r_sc.font.name = "Segoe UI"
    r_sc.font.size = Pt(10)
    r_sc.font.bold = True
    r_sc.font.color.rgb = COLOR_NAVY

    tbl_cond = doc.add_table(rows=1, cols=1)
    tbl_cond.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cond.autofit = False
    c_cond = tbl_cond.rows[0].cells[0]
    c_cond.width = Inches(7.1)
    set_cell_background(c_cond, "FEF3C7") # Warm Amber Tint
    set_cell_margins(c_cond, top=100, bottom=100, left=140, right=140)
    set_cell_borders(c_cond,
                     left={'val': 'single', 'sz': '20', 'color': "D97706"},
                     top={'val': 'single', 'sz': '4', 'color': "FDE68A"},
                     bottom={'val': 'single', 'sz': '4', 'color': "FDE68A"},
                     right={'val': 'single', 'sz': '4', 'color': "FDE68A"})

    p_cond = c_cond.paragraphs[0]
    p_cond.paragraph_format.space_before = Pt(0)
    p_cond.paragraph_format.space_after = Pt(0)
    p_cond.paragraph_format.line_spacing = 1.15
    r_cond_lbl = p_cond.add_run("IMPORTANT TIMELINE CONDITION: ")
    r_cond_lbl.font.name = "Segoe UI"
    r_cond_lbl.font.size = Pt(8.5)
    r_cond_lbl.font.bold = True
    r_cond_lbl.font.color.rgb = RGBColor(146, 64, 14)

    r_cond_txt = p_cond.add_run(
        "The 14-day timeline is fully achievable subject to timely provision of required content, credentials, feedback, "
        "and approvals from the OnPoint Group Limited team. Delays in providing content, access credentials, or approvals "
        "will automatically shift the target launch date by a corresponding duration without liability to DevFlow Technology."
    )
    r_cond_txt.font.name = "Segoe UI"
    r_cond_txt.font.size = Pt(8.5)
    r_cond_txt.font.color.rgb = RGBColor(120, 53, 15)

    # ==========================================
    # 8. COMMUNICATION & PROJECT MANAGEMENT
    # ==========================================
    p_sec6 = doc.add_paragraph()
    p_sec6.paragraph_format.space_before = Pt(16)
    p_sec6.paragraph_format.space_after = Pt(4)
    r_s6_num = p_sec6.add_run("06. ")
    r_s6_num.font.name = "Segoe UI"
    r_s6_num.font.size = Pt(12)
    r_s6_num.font.bold = True
    r_s6_num.font.color.rgb = COLOR_TEAL

    r_s6_txt = p_sec6.add_run("COMMUNICATION & PROJECT MANAGEMENT PROTOCOLS")
    r_s6_txt.font.name = "Segoe UI"
    r_s6_txt.font.size = Pt(12)
    r_s6_txt.font.bold = True
    r_s6_txt.font.color.rgb = COLOR_NAVY

    comm_points = [
        ("Dedicated Communication Channel: ", "A private project communication channel (WhatsApp / Slack / Email) for real-time correspondence and rapid Q&A."),
        ("Daily Asynchronous Progress Updates: ", "Daily end-of-day briefing outlining completed activities and next-day milestones."),
        ("Milestone-Based Synchronous Reviews: ", "Formal video walkthrough sessions scheduled at UI/UX Completion (Day 4) and Final Staging Review (Day 13)."),
        ("Single Point of Contact (SPOC): ", "Prince Gajjar (Founder & CEO, DevFlow Technology) will directly lead technical delivery alongside OnPoint's designated project manager."),
        ("Transparent Escalation Pathway: ", "Direct executive access to resolve any technical, creative, or scheduling blockers within 4 hours.")
    ]

    for cp_title, cp_desc in comm_points:
        p_cp = doc.add_paragraph()
        p_cp.paragraph_format.space_before = Pt(0)
        p_cp.paragraph_format.space_after = Pt(2)
        p_cp.paragraph_format.left_indent = Inches(0.15)
        r_cpt = p_cp.add_run("•  " + cp_title)
        r_cpt.font.name = "Segoe UI"
        r_cpt.font.size = Pt(8.5)
        r_cpt.font.bold = True
        r_cpt.font.color.rgb = COLOR_NAVY
        r_cpd = p_cp.add_run(cp_desc)
        r_cpd.font.name = "Segoe UI"
        r_cpd.font.size = Pt(8.5)
        r_cpd.font.color.rgb = COLOR_BODY

    # ==========================================
    # 9. FINAL DELIVERABLES CHECKLIST
    # ==========================================
    p_sec7 = doc.add_paragraph()
    p_sec7.paragraph_format.space_before = Pt(16)
    p_sec7.paragraph_format.space_after = Pt(4)
    r_s7_num = p_sec7.add_run("07. ")
    r_s7_num.font.name = "Segoe UI"
    r_s7_num.font.size = Pt(12)
    r_s7_num.font.bold = True
    r_s7_num.font.color.rgb = COLOR_TEAL

    r_s7_txt = p_sec7.add_run("FINAL PRODUCTION HANDOVER SPECIFICATION")
    r_s7_txt.font.name = "Segoe UI"
    r_s7_txt.font.size = Pt(12)
    r_s7_txt.font.bold = True
    r_s7_txt.font.color.rgb = COLOR_NAVY

    # 2-column checklist table
    tbl_chk = doc.add_table(rows=5, cols=2)
    tbl_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_chk.autofit = False

    checklist_items = [
        "✓  Modern, Custom-Engineered Responsive UI/UX",
        "✓  Desktop, Tablet & Smartphone Optimization",
        "✓  High-Performance Next.js Frontend Framework",
        "✓  Complete Corporate Content & Media Integration",
        "✓  Technical SEO & Structured Schema Architecture",
        "✓  PageSpeed & Core Web Vitals Optimization",
        "✓  Comprehensive Cross-Browser Quality Assurance",
        "✓  Secured Contact Form & Inquiry Routing",
        "✓  Zero-Downtime Live Production Deployment",
        "✓  Admin Documentation & Final Handover Archive"
    ]

    c_widths = [Inches(3.55), Inches(3.55)]
    idx = 0
    for r in range(5):
        for c in range(2):
            cell = tbl_chk.rows[r].cells[c]
            cell.width = c_widths[c]
            set_cell_background(cell, BG_LIGHT_HEX)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            set_cell_borders(cell,
                             top={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                             bottom={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                             left={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                             right={'val': 'single', 'sz': '4', 'color': BORDER_HEX})
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(checklist_items[idx])
            run.font.name = "Segoe UI"
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = COLOR_NAVY
            idx += 1

    # ==========================================
    # 10. PROJECT SIGN-OFF & AUTHORIZATION
    # ==========================================
    p_sec8 = doc.add_paragraph()
    p_sec8.paragraph_format.space_before = Pt(18)
    p_sec8.paragraph_format.space_after = Pt(4)
    r_s8_num = p_sec8.add_run("08. ")
    r_s8_num.font.name = "Segoe UI"
    r_s8_num.font.size = Pt(12)
    r_s8_num.font.bold = True
    r_s8_num.font.color.rgb = COLOR_TEAL

    r_s8_txt = p_sec8.add_run("PROJECT AUTHORIZATION & ACCEPTANCE")
    r_s8_txt.font.name = "Segoe UI"
    r_s8_txt.font.size = Pt(12)
    r_s8_txt.font.bold = True
    r_s8_txt.font.color.rgb = COLOR_NAVY

    tbl_sig = doc.add_table(rows=1, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sig.autofit = False

    c_sig1 = tbl_sig.rows[0].cells[0]
    c_sig2 = tbl_sig.rows[0].cells[1]
    c_sig1.width = Inches(3.45)
    c_sig2.width = Inches(3.45)

    for cell, entity, rep_name, rep_title in [
        (c_sig1, "FOR DEVFLOW TECHNOLOGY", "Prince Gajjar", "Founder & Chief Executive Officer"),
        (c_sig2, "FOR ONPOINT GROUP LIMITED", "Authorized Representative", "OnPoint Group Limited (Nigeria)")
    ]:
        set_cell_background(cell, BG_LIGHT_HEX)
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        set_cell_borders(cell,
                         top={'val': 'single', 'sz': '6', 'color': NAVY_HEX},
                         bottom={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                         left={'val': 'single', 'sz': '4', 'color': BORDER_HEX},
                         right={'val': 'single', 'sz': '4', 'color': BORDER_HEX})
        
        p_se = cell.paragraphs[0]
        p_se.paragraph_format.space_before = Pt(0)
        p_se.paragraph_format.space_after = Pt(8)
        r_se = p_se.add_run(entity)
        r_se.font.name = "Segoe UI"
        r_se.font.size = Pt(8.5)
        r_se.font.bold = True
        r_se.font.color.rgb = COLOR_NAVY

        p_sline = cell.add_paragraph()
        p_sline.paragraph_format.space_before = Pt(20)
        p_sline.paragraph_format.space_after = Pt(2)
        r_sline = p_sline.add_run("_____________________________________")
        r_sline.font.name = "Segoe UI"
        r_sline.font.size = Pt(8)
        r_sline.font.color.rgb = COLOR_MUTED

        p_sn = cell.add_paragraph()
        p_sn.paragraph_format.space_before = Pt(0)
        p_sn.paragraph_format.space_after = Pt(1)
        r_sn = p_sn.add_run(rep_name)
        r_sn.font.name = "Segoe UI"
        r_sn.font.size = Pt(9)
        r_sn.font.bold = True
        r_sn.font.color.rgb = COLOR_DARK

        p_st = cell.add_paragraph()
        p_st.paragraph_format.space_before = Pt(0)
        p_st.paragraph_format.space_after = Pt(4)
        r_st = p_st.add_run(rep_title)
        r_st.font.name = "Segoe UI"
        r_st.font.size = Pt(8)
        r_st.font.color.rgb = COLOR_MUTED

        p_sd = cell.add_paragraph()
        p_sd.paragraph_format.space_before = Pt(2)
        p_sd.paragraph_format.space_after = Pt(0)
        r_sd = p_sd.add_run("Date: ________________________")
        r_sd.font.name = "Segoe UI"
        r_sd.font.size = Pt(8)
        r_sd.font.color.rgb = COLOR_MUTED

    # Output file
    output_dir = os.path.abspath("timeline")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "ONPOINT_GROUP_Website_Redesign_14_Day_Timeline_DevFlow.docx")
    doc.save(output_path)
    print(f"SUCCESS: Document successfully created at: {output_path}")

if __name__ == "__main__":
    build_document()
